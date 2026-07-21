"""Flatten a .docx VCAA study design into an ordered list of RawBlocks.

Prefer .docx over PDF where available — Word carries real structure
(styles, tables, bold runs) that PDF has to be inferred instead.
"""
from __future__ import annotations

import re

import docx

from .heading_patterns import is_non_glossary_table, looks_like_glossary_row, pattern_level, style_heading_level
from .models import RawBlock


def _heading_level(paragraph) -> int:
    """Heading depth of one paragraph: wording match first, then Word
    style, then "is every run bold in a short paragraph".

    Inputs: paragraph (docx.text.paragraph.Paragraph).
    Outputs: int — level 0-4.
    """
    level = pattern_level(paragraph.text.strip())
    if level:
        return level

    style_level = style_heading_level(paragraph.style.name)
    if style_level:
        return style_level

    runs = [r for r in paragraph.runs if r.text.strip()]
    if runs and all(r.bold for r in runs) and len(paragraph.text) < 60:
        return 4

    return 0


# Matches "VCAA bullet level 2", etc — nested dot points get their own
# style distinct from the top-level bullet.
_SUB_ITEM_STYLE_RE = re.compile(r"level\s*[2-9]", re.I)


def _is_sub_item(paragraph) -> bool:
    """Whether a paragraph is a nested sub-bullet (by style name).

    Inputs: paragraph (docx.text.paragraph.Paragraph).
    Outputs: bool.
    """
    return bool(_SUB_ITEM_STYLE_RE.search(paragraph.style.name))


def parse_docx(path: str) -> list[RawBlock]:
    """Read a .docx file into RawBlocks: every paragraph in reading
    order, then glossary table rows appended at the end.

    Inputs: path (str) — filesystem path to a .docx file.
    Outputs: list[RawBlock].
    """
    document = docx.Document(path)
    blocks: list[RawBlock] = []

    # Every paragraph, skipping blanks. Collapse manual line breaks
    # (Shift+Enter -> literal "\n" in paragraph.text) to a single space.
    for paragraph in document.paragraphs:
        text = re.sub(r"\s+", " ", paragraph.text).strip()
        if not text:
            continue
        blocks.append(
            RawBlock(text=text, level=_heading_level(paragraph), is_sub_item=_is_sub_item(paragraph))
        )

    # Tables: python-docx's paragraph iteration doesn't see table
    # content, so walk document.tables separately. Rows are tagged
    # level=5 for extract_items.py. Tables with a header row that
    # clearly isn't a glossary are skipped entirely; every other
    # table's rows (including the first, in case there's no header at
    # all) go through looks_like_glossary_row.
    for table in document.tables:
        if not table.rows:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        if is_non_glossary_table(header):
            continue

        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2 or not cells[0] or not cells[1]:
                continue
            # A row can contain multiple terms glued together as
            # newline-separated paragraphs in one cell — split back
            # into separate entries when both cells have matching
            # paragraph counts.
            terms = cells[0].split("\n")
            definitions = cells[1].split("\n")
            pairs = (
                zip(terms, definitions)
                if len(terms) == len(definitions) and len(terms) > 1
                else [(cells[0], cells[1])]
            )
            for term, definition in pairs:
                term, definition = term.strip(), definition.strip()
                if term and definition and looks_like_glossary_row(term, definition):
                    blocks.append(RawBlock(text=f"{term}\t{definition}", level=5))

    return blocks
